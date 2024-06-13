import os
import json
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_header_footer(doc):
    # Add header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Security Audit Report"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_paragraph.runs[0]
    run.font.size = Pt(12)
    run.bold = True

    # Add footer with page numbers
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Eloham Caron\nPage "
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.runs[0]
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

def format_firewall_rules(rules):
    formatted_rules = []
    current_rule = {}
    for line in rules:
        if line.strip() == '':
            continue
        if ':' in line:
            key, value = line.split(':', 1)
            current_rule[key.strip()] = value.strip()
        elif line.startswith("Nom de la règle"):
            if current_rule:
                formatted_rules.append(current_rule)
                current_rule = {}
            key, value = line.split(':', 1)
            current_rule[key.strip()] = value.strip()
    if current_rule:
        formatted_rules.append(current_rule)
    return formatted_rules

def add_stylish_cover_page(doc):
    # Page de garde stylisée
    section = doc.sections[0]
    section.page_height = docx.shared.Inches(11)
    section.page_width = docx.shared.Inches(8.5)
    section.top_margin = docx.shared.Inches(1)
    section.bottom_margin = docx.shared.Inches(1)
    section.left_margin = docx.shared.Inches(1)
    section.right_margin = docx.shared.Inches(1)
    
    # Titre principal
    title = doc.add_paragraph()
    title_run = title.add_run("Security Audit Report")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Bleu foncé
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sous-titre
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Comprehensive Security Analysis")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)  # Bleu plus clair
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Informations supplémentaires
    additional_info = doc.add_paragraph()
    additional_info_run = additional_info.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\nAuthor: Eloham Caron")
    additional_info_run.font.size = Pt(16)
    additional_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ligne de séparation
    doc.add_paragraph("\n" * 2)
    line = doc.add_paragraph()
    line_run = line.add_run("─" * 50)
    line_run.font.color.rgb = RGBColor(0, 51, 102)  # Bleu foncé
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Logo ou image (facultatif)
    # doc.add_picture('logo.png', width=docx.shared.Inches(2))

    # Saut de page pour la suite du document
    doc.add_page_break()

def generate_report(results, vulnerabilities, output_file):
    doc = docx.Document()

    # Add header and footer
    add_header_footer(doc)

    # Add stylish cover page
    add_stylish_cover_page(doc)

    # Add table of contents
    doc.add_heading('Table of Contents', level=1)
    add_table_of_contents(doc)
    doc.add_page_break()

    # Section pour les résultats
    for category, data in results.items():
        doc.add_heading(category.replace('_', ' ').title(), level=2)

        if 'firewall' in category.lower():
            formatted_rules = format_firewall_rules(data)
            for rule in formatted_rules:
                for key, value in rule.items():
                    doc.add_paragraph(f"{key}: {value}")
                doc.add_paragraph()  # Add a space between rules
        elif isinstance(data, dict):
            for key, value in data.items():
                if value:  # Vérifier que la valeur n'est pas vide
                    doc.add_paragraph(f"{key}: {value}")
        elif isinstance(data, list):
            for item in data:
                if item:  # Vérifier que l'item n'est pas vide
                    doc.add_paragraph(item)
        else:
            if data:  # Vérifier que la donnée n'est pas vide
                doc.add_paragraph(str(data))
        doc.add_page_break()

    # Section pour les vulnérabilités
    doc.add_heading("Potential Vulnerabilities", level=2)
    for vuln in vulnerabilities:
        doc.add_paragraph(vuln, style='ListBullet')

    # Enregistrement du document
    doc.save(output_file)

def analyze_results(audit_dir):
    results = {}
    for root, _, files in os.walk(audit_dir):
        for file in files:
            file_path = os.path.join(root, file)
            if file.endswith('.json'):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        results[file] = json.load(f)
                except json.JSONDecodeError as e:
                    print(f"Error decoding JSON from {file}: {e}")
                    results[file] = f"Error decoding JSON: {e}"
                except UnicodeDecodeError as e:
                    print(f"Error decoding file {file} with utf-8: {e}")
                    results[file] = f"Error decoding file: {e}"
            elif file.endswith('.txt'):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        results[file] = f.readlines()
                except UnicodeDecodeError:
                    try:
                        with open(file_path, 'r', encoding='latin-1') as f:
                            results[file] = f.readlines()
                    except UnicodeDecodeError as e:
                        print(f"Error decoding file {file} with utf-8 and latin-1: {e}")
                        results[file] = f"Error decoding file: {e}"
    return results

def check_vulnerable_software(installed_software):
    vulnerabilities = []
    # Exemple fictif de vérification des versions vulnérables
    vulnerable_versions = {
        "ExampleSoftware": ["1.0", "1.1"],
    }
    for software in installed_software:
        parts = software.split(',')
        if len(parts) > 1:
            name, version = parts
            name = name.strip()
            version = version.strip()
            if name in vulnerable_versions and version in vulnerable_versions[name]:
                vulnerabilities.append(f"{name} version {version} is vulnerable.")
    return vulnerabilities

def check_vulnerable_services(services):
    vulnerabilities = []
    # Exemple fictif de services vulnérables
    vulnerable_services = ["vulnerable_service_name"]
    for service in services:
        if any(vs in service for vs in vulnerable_services):
            vulnerabilities.append(f"Service {service.strip()} is vulnerable.")
    return vulnerabilities

def main(audit_dir, output_file):
    results = analyze_results(audit_dir)
    vulnerabilities = []

    # Analyser les logiciels installés
    if 'installed_software.txt' in results:
        installed_software = results['installed_software.txt']
        vulnerabilities.extend(check_vulnerable_software(installed_software))

    # Analyser les services
    if 'services_info.txt' in results:
        services = results['services_info.txt']
        vulnerabilities.extend(check_vulnerable_services(services))

    generate_report(results, vulnerabilities, output_file)

if __name__ == "__main__":
    audit_dir = 'Audit'
    output_file = 'Security_Audit_Report.docx'
    main(audit_dir, output_file)
